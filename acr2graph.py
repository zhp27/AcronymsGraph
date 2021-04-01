# -*- codiExtentedExp: utf-8 -*-
"""
Created on Tue Mar 30 20:50:26 2021

@author: raha_
"""

#text canonization
from docx import Document
import re
from nltk import sent_tokenize
import json
import py2neo
from py2neo import Graph, Node, Relationship
uri1="bolt://localhost:7687"


#-----------------------------------------------------------------
#find acronyoms and thier expansion and store them in a dictionary
def expand_abv(Paras):
    #dictionary for those acrs which has an extra word    
    arts = ['a', 'an', 'and', 'the', 'of', 'in', 'with', 'by', 'from']
   #save acrs
    acr={}
    for p in Paras:
        for run in p.runs:
            t=run.text 
            #removiExtentedExp hyphens
            puncts='_-—'
            for sym in puncts:
                t= t.replace(sym,' ')  
                #find acrs (ABCD)
            mainObject=re.findall(r'\([A-Z]*\)', t) 
                         
            words = t.split(' ') 
            #Expand the acr, find ngrams prior to (ABC)
            if mainObject:                                 
                for item in mainObject:                    
                    if item.isupper(): #Doublecheck for acronyms be uppercase                   
                        l=len(item)-1 
                        #print(l)
                        ExtentedExp=''
                        for i, j in enumerate(words):                        
                            if j == item:
                                artno=0 #if there is an Article in the phrase
                                for k in range(1,l):
                                    w=words[i-k]
                                    if w in arts:
                                        artno+=1
                                    ExtentedExp+=w
                                    ExtentedExp+=' '
                                #of the ,... management. 
                                #expand the phrase consideriExtentedExp articles 
                                if artno>0:
                                    for c in range(artno):                                    
                                        w=words[i-(k+c+1)]
                                        ExtentedExp+=w
                                        ExtentedExp+=' '
                                #reverse the phrase
                                sExtentedExp=ExtentedExp.split()        
                                l1= len(sExtentedExp) 
                                rExtentedExp=''
                                for i in reversed(range(l1)):
                                    rExtentedExp+=sExtentedExp[i]
                                    rExtentedExp+=' '                               
                                    #add to dictionary
                                
                                acr[item]=rExtentedExp
    
    facr={}
    #print(json.dumps(acr, indent = 4))
    for item, rExtentedExp in acr.items():
        i=item.replace('(','')
        i=i.replace(')','')
        facr[i]=rExtentedExp
        
    #print(json.dumps(facr, indent = 4))
    return facr

    
#expand the acrs
def acrs(text,da): 
    if da:
        #convert dict of acrss to list
        acronyms=[]
        temp=[]
        for key, value in da.items():
            temp = [key,value]
            acronyms.append(temp) 
         
        
        pattern = '|'.join('(%s)' % re.escape(match) for match, replacement in acronyms) 
        substitutions = [match for replacement, match in acronyms]
        replace = lambda m: substitutions[m.lastindex - 1]
        ft=re.sub(pattern, replace, text)
    
        return ft,acronyms
    else: return text, []

#remove the (ACR) to prevent duplicate words
def remove_first(Text):
    
    mainObject=re.findall(r'\([A-Z]*\)', Text) 
        #print(mo)
    if mainObject:
        for item in mainObject: 
            Text=Text.replace(item,'')
    
    return(Text)

def acr_grapgh(ns,nd,l):   
    
    for i in range(l):
        
        s=ns[i]
        d=nd[i]
        n1 = Node("acronym",name=s)
        n2 = Node("expansion",name=d)
        e='has acronym'
        
        rel = Relationship(n2,e,n1) #n2-RelationshipType->n1
        graph = Graph(uri1, user='neo4j', password='123456')
        tx = graph.begin()
        tx.merge(n1,"acronym","name") #node,label,primary key
        tx.merge(n2,"expansion","name") #node,label,pirmary key
        tx.merge(rel)
        tx.commit()

#Main code
doc = Document('acr.docx')
Paras = doc.paragraphs
#create a list of acronyms
DictAcr=expand_abv(Paras)
Paras=doc.paragraphs
for p in Paras:    
    for run in p.runs:        
        T=run.text 
        #remove ACR instance from text eg genetic alg (ga)> genetic alg
        T=remove_first(T)
        #resolve and expand all the Acronyms
        [T,dic]=acrs(T,DictAcr)   
#start node, End node creation
LAcr=len(dic)
ns=[]
nd=[]
for i in range(LAcr):
    ns.append(dic[i][0])
    nd.append(dic[i][1])

acr_grapgh(ns, nd, LAcr)
